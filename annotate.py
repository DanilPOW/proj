import fitz
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import tempfile
import os
import zipfile
import xml.etree.ElementTree as ET
import re
from lxml import etree
from fuzzywuzzy import fuzz

def replace_cyrillic_with_latin_extended(text):
    """
    РАСШИРЕННАЯ замена кириллических символов на латинские
    """
    cyrillic_to_latin = {
        # Основные буквы
        'А': 'A', 'а': 'a',
        'В': 'V', 'в': 'v', 
        'Е': 'E', 'е': 'e',
        'К': 'K', 'к': 'k',
        'М': 'M', 'м': 'm',
        'Н': 'H', 'н': 'h',
        'О': 'O', 'о': 'o',
        'Р': 'P', 'р': 'p',
        'С': 'C', 'с': 'c',
        'Т': 'T', 'т': 't',
        'У': 'U', 'у': 'u',
        'Х': 'X', 'х': 'x',
        'Ѕ': 'S', 'ѕ': 's',
        
        # Специальные математические символы
        '𝐸': 'E', '𝑚': 'm', '𝑐': 'c',
        '𝑈': 'U', '𝑅': 'R', '𝐼': 'I',
        
        # Единицы измерения
        'кОм': 'kOm', 'Ом': 'Om',
        'кг': 'kg', 'м/с': 'm/s',
        'Гц': 'Hz', 'Вт': 'W',
    }
    
    result = text
    for cyr, lat in cyrillic_to_latin.items():
        result = result.replace(cyr, lat)
    
    return result

def separate_merged_variables(text):
    """
    Разделяет слитые переменные типа VR -> V/R, UI -> U/I
    """
    # Паттерны для разделения слитых переменных
    patterns = [
        # Две заглавные буквы подряд в конце или перед =
        (r'\b([A-Z])([A-Z])\b(?=\s*=|\s*$)', r'\1/\2'),
        
        # Специальные случаи для физических формул
        (r'\bUI\b', 'U/I'),
        (r'\bVR\b', 'V/R'), 
        (r'\bIR\b', 'I/R'),
        (r'\bPV\b', 'P/V'),
        (r'\bFm\b', 'F/m'),
        
        # Переменная + цифра (степень)
        (r'\b([A-Za-z])\s*(\d)\b', r'\1^\2'),
    ]
    
    result = text
    for pattern, replacement in patterns:
        old_result = result
        result = re.sub(pattern, replacement, result)
        if result != old_result:
            print(f"    🔧 РАЗДЕЛЕНИЕ ПЕРЕМЕННЫХ: '{old_result}' -> '{result}'")
    
    return result

def fix_pdf_formula_structure(text):
    """
    ИСПРАВЛЕНИЕ структуры PDF формул с учетом всех артефактов
    """
    if not text:
        return text
    
    original_text = text
    
    # 1. Заменяем кириллические символы
    text = replace_cyrillic_with_latin_extended(text)
    
    # 2. Исправляем Unicode символы
    unicode_fixes = {
        '∙': '*', '×': '*', '⋅': '*', '·': '*',
        '−': '-', '–': '-', '—': '-',
        '÷': '/', '²': '^2', '³': '^3',
        '===': '=', '==': '=',
    }
    
    for old, new in unicode_fixes.items():
        text = text.replace(old, new)
    
    # 3. СПЕЦИАЛЬНАЯ ОБРАБОТКА: цифра в начале + переменные
    # "2𝐸= 𝑚𝑐" -> "E=mc^2"
    match = re.match(r'^(\d+)\s*([A-Za-z])\s*=\s*(.+)$', text)
    if match:
        digit, variable, rest = match.groups()
        if digit in ['2', '3', '4', '5']:
            # Ищем последнюю переменную для добавления степени
            rest_parts = rest.split()
            if rest_parts:
                last_part = rest_parts[-1]
                if len(last_part) == 1 and last_part.isalpha():
                    rest_parts[-1] = f"{last_part}^{digit}"
                    rest = ' '.join(rest_parts)
                elif 'c' in rest.lower():
                    rest = re.sub(r'\bc\b', f'c^{digit}', rest, flags=re.IGNORECASE)
            
            text = f"{variable}={rest}"
            print(f"    🔧 ИСПРАВЛЕН ПОРЯДОК СТЕПЕНИ: '{original_text}' -> '{text}'")
    
    # 4. Разделяем слитые переменные
    text = separate_merged_variables(text)
    
    # 5. УМНАЯ обрезка по единицам измерения
    # Ищем первое появление единиц, но только после знака =
    if '=' in text:
        eq_pos = text.find('=')
        after_eq = text[eq_pos:]
        
        units_pattern = r'\b(kOm|Om|Ω|V|A|kg|m/s|Hz|W|кОм|Ом|В|А|кг|м/с|Гц|Вт)\b'
        unit_match = re.search(units_pattern, after_eq, flags=re.IGNORECASE)
        
        if unit_match:
            # Обрезаем до единицы измерения
            cut_pos = eq_pos + unit_match.start()
            before_unit = text[:cut_pos].strip()
            
            # Убираем запятые и лишние символы в конце
            before_unit = re.sub(r'[,;.\s]+$', '', before_unit)
            
            if len(before_unit) >= 3 and '=' in before_unit:
                text = before_unit
                print(f"    ✂️ ОБРЕЗКА ПО ЕДИНИЦАМ: '{original_text}' -> '{text}'")
    
    # 6. Убираем лишние пробелы и нормализуем
    text = re.sub(r'\s+', ' ', text.strip())
    text = re.sub(r'\s*([=+\-*/^])\s*', r'\1', text)  # Убираем пробелы вокруг операторов
    
    if text != original_text:
        print(f"    🧹 PDF СТРУКТУРНАЯ ОЧИСТКА: '{original_text}' -> '{text}'")
    
    return text

def fix_docx_formula_structure(text):
    """
    ИСПРАВЛЕНИЕ структуры DOCX формул
    """
    if not text:
        return text
    
    original_text = text
    
    # 1. Заменяем кириллические символы
    text = replace_cyrillic_with_latin_extended(text)
    
    # 2. Убираем OMML артефакты и дублирование
    # "E=mc2c^2" -> "E=mc^2"
    text = re.sub(r'([a-z])(\d)([a-z])\^(\d)', r'\1\3^\2', text)  # mc2c^2 -> mc^2
    text = re.sub(r'([a-z])\^?(\d)\1\^(\d)', r'\1^\2', text)      # c2c^2 -> c^2
    
    # 3. Восстанавливаем дроби из слитых переменных
    # "UI" -> "U/I", но только если это не часть большего выражения
    restoration_patterns = [
        (r'\bUI\b(?!=)', 'U/I'),  # UI не перед =
        (r'\bVR\b(?!=)', 'V/R'),
        (r'\bIR\b(?!=)', 'I/R'),
        (r'\bPV\b(?!=)', 'P/V'),
    ]
    
    for pattern, replacement in restoration_patterns:
        old_text = text
        text = re.sub(pattern, replacement, text)
        if text != old_text:
            print(f"    🔧 ВОССТАНОВЛЕНА ДРОБЬ DOCX: '{old_text}' -> '{text}'")
    
    # 4. Обрезаем пояснительные хвосты РАНЬШЕ
    split_patterns = [
        r'[;,]\s*(?=[А-ЯA-Z])',  # ; где, , При
        r'\s+где\s+',
        r'\s+при\s+',
        r'\s+для\s+',
        r'$$[^)]*$$$$[^)]*$$',   # Убираем двойные скобки (U)/(I)(2 В)/(1∙10-3 А)
    ]
    
    for pattern in split_patterns:
        parts = re.split(pattern, text, maxsplit=1)
        if len(parts) > 1 and '=' in parts[0]:
            text = parts[0].strip()
            break
    
    # 5. Обрезаем по единицам измерения
    units_pattern = r'\b(kOm|Om|Ω|V|A|kg|m/s|Hz|W)\b'
    unit_matches = list(re.finditer(units_pattern, text, flags=re.IGNORECASE))
    
    if unit_matches:
        first_unit = unit_matches[0]
        before_unit = text[:first_unit.start()].strip()
        
        if '=' in before_unit and len(before_unit) >= 3:
            text = before_unit
            print(f"    ✂️ DOCX ОБРЕЗКА ПО ЕДИНИЦАМ: '{original_text}' -> '{text}'")
    
    # 6. Убираем лишние символы и нормализуем
    text = re.sub(r'[,;.\s]+$', '', text)
    text = re.sub(r'=+', '=', text)
    text = re.sub(r'\s+', ' ', text.strip())
    text = re.sub(r'\s*([=+\-*/^])\s*', r'\1', text)
    
    if text != original_text:
        print(f"    🧹 DOCX СТРУКТУРНАЯ ОЧИСТКА: '{original_text}' -> '{text}'")
    
    return text

def tokenize_formula(text):
    """
    Разбивает формулу на математические токены для сравнения
    """
    if not text:
        return []
    
    # Извлекаем токены: переменные, числа, операторы
    tokens = re.findall(r'[A-Za-z]+|\d+(?:\^\d+)?|\d*\.\d+|[=+\-*/^<>≤≥≠≈]', text)
    
    # Нормализуем токены
    normalized_tokens = []
    for token in tokens:
        token = token.lower()
        # Стандартизируем операторы
        if token in ['×', '∙', '·', '⋅']:
            token = '*'
        elif token in ['−', '–', '—']:
            token = '-'
        elif token in ['÷']:
            token = '/'
        
        normalized_tokens.append(token)
    
    return normalized_tokens

def calculate_token_similarity(tokens1, tokens2):
    """
    Вычисляет схожесть на основе токенов (Jaccard distance)
    """
    if not tokens1 or not tokens2:
        return 0.0
    
    set1 = set(tokens1)
    set2 = set(tokens2)
    
    intersection = set1 & set2
    union = set1 | set2
    
    if not union:
        return 0.0
    
    jaccard = len(intersection) / len(union)
    
    # Бонус за одинаковое количество токенов
    length_bonus = 1 - abs(len(tokens1) - len(tokens2)) / max(len(tokens1), len(tokens2))
    
    # Бонус за правильный порядок ключевых операторов
    order_bonus = 0
    if '=' in tokens1 and '=' in tokens2:
        try:
            eq_pos1 = tokens1.index('=')
            eq_pos2 = tokens2.index('=')
            # Проверяем, что переменные до = похожи
            before_eq1 = set(tokens1[:eq_pos1])
            before_eq2 = set(tokens2[:eq_pos2])
            if before_eq1 & before_eq2:
                order_bonus = 0.1
        except ValueError:
            pass
    
    total_score = jaccard + length_bonus * 0.2 + order_bonus
    
    return min(total_score, 1.0)

def is_mathematical_formula_strict_v3(text):
    """
    УЖЕСТОЧЕННАЯ функция определения математических формул v3
    """
    if not text or len(text.strip()) < 2:
        return False
    
    text = text.strip()
    text = fix_pdf_formula_structure(text)
    
    # ЖЕСТКИЕ ИСКЛЮЧЕНИЯ
    strict_exclusions = [
        # Библиографические записи
        r'\b(ISBN|URL|ГОСТ|МПК|DOI|http|www\.)\b',
        r'\b(Москва|СПб|Санкт-Петербург|Пенза|Екатеринбург)\b',
        r'\d{4}\s*г\.',
        r'[А-ЯЁ][а-яё]+,\s*[А-ЯЁ]\.\s*[А-ЯЁ]\.',  # Фамилия, И. О.
        
        # Заголовки и разделы
        r'^\d+(\.\d+)*\s+[А-ЯЁ]',
        r'^[А-ЯЁ][А-ЯЁ\s]+\d*$',
        r'^\d+\s*$',
        r'^[А-ЯЁ][А-ЯЁ\s]{10,}$',
        
        # Нормативные фразы
        r'должно быть не более',
        r'должно быть не менее',
        r'в соответствии с',
        r'согласно ГОСТ',
        
        # Списки и перечисления (УСИЛЕНО)
        r'^\s*[–—]\s*[а-яё]',  # – применять
        r'^\d+\)\s*[а-яё]',
        r'^[а-яё]\)\s*',
        
        # Пояснения и расшифровки (НОВОЕ)
        r'^[a-zA-Zа-яё]\s*[–—]\s*[а-яё]',  # c – скорость
        r'[а-яё]{5,}\s*[а-яё]{5,}',        # длинные русские слова
        
        # Страницы и ссылки
        r'стр\.\s*\d+',
        r'страница\s*\d+',
        r'рисун[ок|ке]\s*\d+',
        r'таблиц[а|е]\s*\d+',
        
        # Очень длинные строки
        r'^.{150,}$',
        
        # Простые числовые последовательности
        r'^\d+(\s*[",]\s*\d+)*$',  # "1 1", "4 2"
    ]
    
    for pattern in strict_exclusions:
        if re.search(pattern, text, re.IGNORECASE):
            print(f"    ❌ ИСКЛЮЧЕНО v3 (паттерн '{pattern[:30]}...'): '{text[:50]}...'")
            return False
    
    # ПОЗИТИВНЫЕ ПРИЗНАКИ - должно быть несколько
    math_indicators = [
        r'[=≠≈≤≥<>]',                    # Основные операторы
        r'[A-Za-z]\s*/\s*[A-Za-z]',     # Дроби переменных
        r'[A-Za-z]\^[0-9]',             # Степени
        r'[A-Za-z][²³⁴⁵⁶⁷⁸⁹]',         # Unicode степени
        r'\b[EIURPVF]\s*=',             # Физические переменные
        r'\d+\s*[+\-*/]\s*\d+',        # Числовые выражения
        r'(sin|cos|tan|log|ln|exp|sqrt)\s*\(',  # Функции
        r'[αβγδεζηθικλμνξοπρστυφχψω]',  # Греческие буквы
        r'10\^[+-]?\d+',                # Научная нотация
    ]
    
    indicator_count = sum(1 for pattern in math_indicators if re.search(pattern, text, re.IGNORECASE))
    
    # Требуем минимум 2 математических признака для коротких строк
    min_indicators = 2 if len(text) < 30 else 1
    
    if indicator_count >= min_indicators:
        print(f"    ✅ ПРИНЯТО v3 ({indicator_count} мат. признаков): '{text[:50]}...'")
        return True
    
    print(f"    ❌ ОТКЛОНЕНО v3 ({indicator_count} признаков < {min_indicators}): '{text[:50]}...'")
    return False

def improved_fuzzy_matching_v3(text1, text2):
    """
    КАРДИНАЛЬНО УЛУЧШЕННОЕ сопоставление v3 с токенизацией
    """
    if not text1 or not text2:
        return 0.0
    
    # Структурная очистка
    clean1 = fix_pdf_formula_structure(text1) if any(c in text1 for c in '𝑈𝑅𝐼В') else fix_docx_formula_structure(text1)
    clean2 = fix_docx_formula_structure(text2)
    
    print(f"      Очищенные: '{clean1}' vs '{clean2}'")
    
    # Токенизация
    tokens1 = tokenize_formula(clean1)
    tokens2 = tokenize_formula(clean2)
    
    print(f"      Токены: {tokens1} vs {tokens2}")
    
    # Точное совпадение токенов
    if tokens1 == tokens2 and tokens1:
        return 100.0
    
    # Схожесть по токенам (основной метод)
    token_score = calculate_token_similarity(tokens1, tokens2) * 100
    
    # Дополнительные методы для подстраховки
    scores = [token_score]
    
    # Обычные fuzzy методы
    scores.append(fuzz.ratio(clean1, clean2))
    scores.append(fuzz.partial_ratio(clean1, clean2))
    scores.append(fuzz.token_set_ratio(clean1, clean2))
    scores.append(fuzz.token_sort_ratio(clean1, clean2))
    scores.append(fuzz.WRatio(clean1, clean2))
    
    # Структурная схожесть
    struct_score = variable_overlap_score(clean1, clean2)
    scores.append(struct_score)
    
    best_score = max(scores) / 100.0 if max(scores) > 100 else max(scores) / 100.0
    
    print(f"      Scores: token={token_score:.1f}, ratio={scores[1]}, partial={scores[2]}, token_set={scores[3]}, token_sort={scores[4]}, WRatio={scores[5]}, struct={scores[6]:.1f}")
    print(f"      BEST: {best_score:.2f}")
    
    return best_score

def variable_overlap_score(text1, text2):
    """
    Вычисляет схожесть по переменным и операторам
    """
    if not text1 or not text2:
        return 0.0
    
    # Извлекаем переменные (одиночные буквы)
    vars1 = set(re.findall(r'\b[A-Za-z]\b', text1))
    vars2 = set(re.findall(r'\b[A-Za-z]\b', text2))
    
    # Извлекаем операторы
    ops1 = set(re.findall(r'[=+\-*/^<>≤≥≠≈]', text1))
    ops2 = set(re.findall(r'[=+\-*/^<>≤≥≠≈]', text2))
    
    # Вычисляем пересечения
    var_overlap = len(vars1 & vars2) / max(len(vars1 | vars2), 1)
    op_overlap = len(ops1 & ops2) / max(len(ops1 | ops2), 1)
    
    # Взвешенная сумма
    return (var_overlap * 0.7 + op_overlap * 0.3) * 100

def normalize_formula_text_advanced(text):
    """
    ПРОДВИНУТАЯ нормализация с исправлением PDF артефактов
    """
    if not text:
        return ""
    
    # Сначала структурная очистка
    text = fix_pdf_formula_structure(text) if any(c in text for c in '𝑈𝑅𝐼В') else fix_docx_formula_structure(text)
    
    # Базовая нормализация
    normalized = normalize_formula_text(text)
    
    return normalized

def normalize_formula_text(text):
    """
    БАЗОВАЯ нормализация текста формулы
    """
    if not text:
        return ""
    
    # Приводим к нижнему регистру
    normalized = text.lower()
    
    # Убираем все пробелы
    normalized = re.sub(r'\s+', '', normalized)
    
    # Стандартизируем математические символы
    symbol_replacements = {
        '×': '*', '÷': '/', '−': '-', '–': '-', '—': '-',
        '∙': '*', '·': '*', '∗': '*', '⋅': '*',
        '≈': '~', '≠': '!=', '≤': '<=', '≥': '>=',
        '∞': 'inf', '∑': 'sum', '∫': 'int', '∂': 'd',
        '∆': 'delta', '∇': 'nabla', '±': '+-', '∓': '-+',
        '√': 'sqrt', '∛': 'cbrt', '∜': 'qrt',
        '²': '^2', '³': '^3', '⁴': '^4', '⁵': '^5',
        '⁶': '^6', '⁷': '^7', '⁸': '^8', '⁹': '^9',
        '⁰': '^0', '¹': '^1',
        '½': '1/2', '⅓': '1/3', '¼': '1/4', '¾': '3/4',
        '⅛': '1/8', '⅜': '3/8', '⅝': '5/8', '⅞': '7/8'
    }
    
    for old_symbol, new_symbol in symbol_replacements.items():
        normalized = normalized.replace(old_symbol, new_symbol)
    
    # Убираем специальные символы, оставляя только буквы, цифры и основные математические операторы
    normalized = re.sub(r'[^\w\+\-\*\/\=$$$$\[\]\{\}\^\<\>\!\~]', '', normalized)
    
    # Убираем повторяющиеся символы
    normalized = re.sub(r'(.)\1+', r'\1', normalized)
    
    return normalized

def extract_omml_formulas_with_lxml(docx_path):
    """
    ИСПРАВЛЕННАЯ функция прямого извлечения формул через lxml и парсинг OMML
    """
    formulas_info = []
    
    try:
        print("=== ИЗВЛЕЧЕНИЕ OMML ФОРМУЛ ЧЕРЕЗ LXML ===")
        
        # Открываем DOCX как ZIP архив
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Читаем document.xml
            try:
                document_xml = docx_zip.read('word/document.xml')
                print("✓ Успешно прочитан word/document.xml")
            except KeyError:
                print("❌ Не найден word/document.xml")
                return formulas_info
            
            # Парсим XML с помощью lxml
            try:
                root = etree.fromstring(document_xml)
                print("✓ XML успешно распарсен")
            except Exception as e:
                print(f"❌ Ошибка парсинга XML: {e}")
                return formulas_info
            
            # Определяем пространства имен
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
                'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006'
            }
            
            # Ищем все математические объекты
            math_elements = root.xpath('.//m:oMath | .//m:oMathPara', namespaces=namespaces)
            print(f"Найдено {len(math_elements)} математических объектов OMML")
            
            # Обрабатываем каждый математический объект
            for idx, math_elem in enumerate(math_elements):
                try:
                    # Извлекаем текст из математического объекта
                    formula_text = extract_text_from_omml(math_elem, namespaces)
                    
                    if formula_text and formula_text.strip():
                        # СТРОГАЯ ПРОВЕРКА: это действительно формула?
                        if not is_mathematical_formula_strict_v3(formula_text.strip()):
                            print(f"  ❌ OMML объект {idx+1} не является формулой: '{formula_text.strip()}'")
                            continue
                        
                        # Находим родительский параграф
                        paragraph_elem = math_elem
                        while paragraph_elem is not None and paragraph_elem.tag != f"{{{namespaces['w']}}}p":
                            paragraph_elem = paragraph_elem.getparent()
                        
                        # Определяем позицию в документе (приблизительно)
                        all_paragraphs = root.xpath('.//w:p', namespaces=namespaces)
                        paragraph_idx = 0
                        if paragraph_elem is not None:
                            try:
                                paragraph_idx = all_paragraphs.index(paragraph_elem)
                            except ValueError:
                                paragraph_idx = idx  # Fallback
                        
                        # Получаем контекст
                        context_before, context_after = get_formula_context_from_xml(
                            root, paragraph_elem, namespaces, paragraph_idx
                        )
                        
                        # Проверяем центрированность
                        is_centered = check_formula_alignment_in_xml(paragraph_elem, namespaces)
                        
                        formulas_info.append({
                            'paragraph_idx': paragraph_idx,
                            'text': formula_text.strip(),
                            'normalized_text': normalize_formula_text_advanced(formula_text),
                            'context_before': context_before,
                            'context_after': context_after,
                            'has_math_object': True,
                            'is_centered': is_centered,
                            'position_in_doc': paragraph_idx,
                            'type': 'omml_extracted',
                            'confidence': 'high',
                            'extraction_method': 'lxml_omml'
                        })
                        
                        print(f"  ✅ OMML формула {idx+1}: '{formula_text.strip()}'")
                        print(f"      Нормализованная: '{normalize_formula_text_advanced(formula_text)}'")
                        print(f"      Параграф: {paragraph_idx}, Центрирована: {is_centered}")
                    
                except Exception as e:
                    print(f"  ❌ Ошибка обработки OMML формулы {idx+1}: {e}")
                    continue
    
    except Exception as e:
        print(f"❌ Общая ошибка извлечения OMML: {e}")
    
    print(f"=== ИЗВЛЕЧЕНО {len(formulas_info)} НАСТОЯЩИХ OMML ФОРМУЛ ===")
    return formulas_info

def extract_text_from_omml(math_elem, namespaces):
    """
    Извлекает текст из OMML математического элемента
    """
    try:
        # Ищем все текстовые элементы в математическом объекте
        text_elements = math_elem.xpath('.//m:t', namespaces=namespaces)
        
        formula_parts = []
        for text_elem in text_elements:
            if text_elem.text:
                formula_parts.append(text_elem.text)
        
        # Также ищем обычные текстовые элементы внутри математики
        w_text_elements = math_elem.xpath('.//w:t', namespaces=namespaces)
        for text_elem in w_text_elements:
            if text_elem.text:
                formula_parts.append(text_elem.text)
        
        # Объединяем части формулы
        formula_text = ''.join(formula_parts)
        
        # Дополнительная обработка для извлечения структурных элементов
        # Ищем дроби
        fractions = math_elem.xpath('.//m:f', namespaces=namespaces)
        for frac in fractions:
            num_text = ''.join([t.text or '' for t in frac.xpath('.//m:num//m:t', namespaces=namespaces)])
            den_text = ''.join([t.text or '' for t in frac.xpath('.//m:den//m:t', namespaces=namespaces)])
            if num_text and den_text:
                formula_text += f"({num_text})/({den_text})"
        
        # Ищем степени
        superscripts = math_elem.xpath('.//m:sSup', namespaces=namespaces)
        for sup in superscripts:
            base_text = ''.join([t.text or '' for t in sup.xpath('.//m:e//m:t', namespaces=namespaces)])
            sup_text = ''.join([t.text or '' for t in sup.xpath('.//m:sup//m:t', namespaces=namespaces)])
            if base_text and sup_text:
                formula_text += f"{base_text}^{sup_text}"
        
        return formula_text
        
    except Exception as e:
        print(f"Ошибка извлечения текста из OMML: {e}")
        return ""

def get_formula_context_from_xml(root, paragraph_elem, namespaces, paragraph_idx):
    """
    Получает контекст формулы из XML
    """
    try:
        all_paragraphs = root.xpath('.//w:p', namespaces=namespaces)
        
        context_before = ""
        context_after = ""
        
        # Контекст до формулы
        for i in range(max(0, paragraph_idx - 2), paragraph_idx):
            if i < len(all_paragraphs):
                para_text = get_paragraph_text_from_xml(all_paragraphs[i], namespaces)
                context_before += para_text + " "
        
        # Контекст после формулы
        for i in range(paragraph_idx + 1, min(len(all_paragraphs), paragraph_idx + 3)):
            if i < len(all_paragraphs):
                para_text = get_paragraph_text_from_xml(all_paragraphs[i], namespaces)
                context_after += para_text + " "
        
        return context_before.strip()[-100:], context_after.strip()[:100]
        
    except Exception as e:
        print(f"Ошибка получения контекста: {e}")
        return "", ""

def get_paragraph_text_from_xml(paragraph_elem, namespaces):
    """
    Извлекает текст из параграфа XML
    """
    try:
        text_elements = paragraph_elem.xpath('.//w:t', namespaces=namespaces)
        return ''.join([elem.text or '' for elem in text_elements])
    except:
        return ""

def check_formula_alignment_in_xml(paragraph_elem, namespaces):
    """
    Проверяет центрированность формулы в XML
    """
    try:
        if paragraph_elem is None:
            return False
        
        # Ищем настройки выравнивания параграфа
        jc_elements = paragraph_elem.xpath('.//w:jc', namespaces=namespaces)
        for jc in jc_elements:
            val = jc.get(f"{{{namespaces['w']}}}val")
            if val == "center":
                return True
        
        return False
        
    except Exception as e:
        print(f"Ошибка проверки выравнивания: {e}")
        return False

def find_docx_formulas_with_positions(docx_path):
    """
    ФИНАЛЬНАЯ ИСПРАВЛЕННАЯ функция поиска формул в DOCX v3
    """
    print("=== ПОИСК ФОРМУЛ В DOCX (ФИНАЛЬНАЯ ВЕРСИЯ v3) ===")
    
    # Метод 1: Извлечение через lxml и OMML (приоритетный)
    omml_formulas = extract_omml_formulas_with_lxml(docx_path)
    
    # Метод 2: Дополнительное извлечение через python-docx с УЖЕСТОЧЕННОЙ фильтрацией v3
    try:
        doc = Document(docx_path)
        docx_formulas = []
        
        print("\n=== РЕЗЕРВНОЕ ИЗВЛЕЧЕНИЕ ЧЕРЕЗ PYTHON-DOCX (v3) ===")
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            
            if not para_text:
                continue
            
            print(f"  Проверяем параграф {para_idx}: '{para_text[:50]}{'...' if len(para_text) > 50 else ''}'")
            
            # Проверяем, есть ли уже эта формула в OMML результатах
            found_in_omml = False
            for omml_formula in omml_formulas:
                if (omml_formula['paragraph_idx'] == para_idx or 
                    abs(omml_formula['paragraph_idx'] - para_idx) <= 1):
                    found_in_omml = True
                    print(f"    ⚠️ Уже найдено в OMML")
                    break
            
            if not found_in_omml and is_mathematical_formula_strict_v3(para_text):
                # Получаем контекст
                context_before = ""
                context_after = ""
                
                for i in range(max(0, para_idx - 2), para_idx):
                    context_before += doc.paragraphs[i].text + " "
                
                for i in range(para_idx + 1, min(len(doc.paragraphs), para_idx + 3)):
                    context_after += doc.paragraphs[i].text + " "
                
                # Проверяем центрированность
                is_centered = False
                if paragraph.paragraph_format.alignment is not None:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    is_centered = paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
                
                docx_formulas.append({
                    'paragraph_idx': para_idx,
                    'text': para_text,
                    'normalized_text': normalize_formula_text_advanced(para_text),
                    'context_before': context_before.strip()[-100:],
                    'context_after': context_after.strip()[:100],
                    'has_math_object': False,
                    'is_centered': is_centered,
                    'position_in_doc': para_idx,
                    'type': 'text_formula',
                    'confidence': 'medium',
                    'extraction_method': 'python_docx'
                })
                
                print(f"    ✅ ДОПОЛНИТЕЛЬНАЯ ФОРМУЛА: параграф {para_idx}")
                print(f"        Текст: '{para_text}'")
        
        # Объединяем результаты (OMML имеет приоритет)
        all_formulas = omml_formulas + docx_formulas
        
    except Exception as e:
        print(f"❌ Ошибка резервного извлечения: {e}")
        all_formulas = omml_formulas
    
    print(f"\n=== ИТОГО НАЙДЕНО {len(all_formulas)} НАСТОЯЩИХ ФОРМУЛ В DOCX ===")
    print(f"OMML формул: {len(omml_formulas)}")
    print(f"Текстовых формул: {len(all_formulas) - len(omml_formulas)}")
    
    # Выводим все найденные формулы для проверки
    for i, formula in enumerate(all_formulas, 1):
        print(f"  📐 Формула {i}: '{formula['text']}' (метод: {formula['extraction_method']}, поз: {formula['position_in_doc']})")
    
    return all_formulas

def match_pdf_formulas_to_docx(pdf_formulas, docx_formulas):
    """
    ФИНАЛЬНОЕ ИСПРАВЛЕННОЕ сопоставление формул v3
    """
    matching = {}
    
    if not pdf_formulas or not docx_formulas:
        print(f"❌ Недостаточно данных для сопоставления: PDF={len(pdf_formulas or [])}, DOCX={len(docx_formulas or [])}")
        return matching
    
    print(f"=== ФИНАЛЬНОЕ СОПОСТАВЛЕНИЕ ФОРМУЛ v3 ===")
    print(f"PDF формул: {len(pdf_formulas)}")
    print(f"DOCX формул: {len(docx_formulas)}")
    
    # Фильтруем PDF формулы - убираем явно неправильные
    filtered_pdf_formulas = []
    for i, pdf_formula in enumerate(pdf_formulas):
        pdf_text = pdf_formula.get('text', '').strip()
        if is_mathematical_formula_strict_v3(pdf_text):
            filtered_pdf_formulas.append((i, pdf_formula))
        else:
            print(f"❌ PDF формула {i+1} отфильтрована: '{pdf_text}'")
    
    print(f"После фильтрации PDF: {len(filtered_pdf_formulas)} из {len(pdf_formulas)}")
    
    # Нормализуем все формулы для сравнения
    normalized_pdf = []
    normalized_docx = []
    
    for original_idx, pdf_formula in filtered_pdf_formulas:
        cleaned_text = fix_pdf_formula_structure(pdf_formula.get('text', ''))
        normalized_text = normalize_formula_text_advanced(cleaned_text)
        normalized_pdf.append({
            'original_index': original_idx,
            'index': len(normalized_pdf),
            'original_text': pdf_formula.get('text', ''),
            'cleaned_text': cleaned_text,
            'normalized_text': normalized_text,
            'page': pdf_formula.get('page', 1),
            'context': (pdf_formula.get('context_before', '') + ' ' + 
                       pdf_formula.get('context_after', '')).strip()
        })
        print(f"PDF {original_idx+1}: '{pdf_formula.get('text', '')}' -> '{cleaned_text}' -> '{normalized_text}'")
    
    for i, docx_formula in enumerate(docx_formulas):
        cleaned_text = fix_docx_formula_structure(docx_formula.get('text', ''))
        normalized_text = normalize_formula_text_advanced(cleaned_text)
        normalized_docx.append({
            'index': i,
            'original_text': docx_formula.get('text', ''),
            'cleaned_text': cleaned_text,
            'normalized_text': normalized_text,
            'position': docx_formula.get('position_in_doc', i),
            'context': (docx_formula.get('context_before', '') + ' ' + 
                       docx_formula.get('context_after', '')).strip()
        })
        print(f"DOCX {i+1}: '{docx_formula.get('text', '')}' -> '{cleaned_text}' -> '{normalized_text}'")
    
    used_docx = set()
    match_log = []
    
    # МЕТОД 1: Точное сопоставление по нормализованному тексту
    print("\n🎯 МЕТОД 1: Точное сопоставление")
    for pdf_item in normalized_pdf:
        if pdf_item['original_index'] in matching:
            continue
            
        for docx_item in normalized_docx:
            if docx_item['index'] in used_docx:
                continue
                
            if (pdf_item['normalized_text'] == docx_item['normalized_text'] and 
                pdf_item['normalized_text'] and 
                len(pdf_item['normalized_text']) > 2):
                
                matching[pdf_item['original_index']] = docx_item['index']
                used_docx.add(docx_item['index'])
                
                log_entry = {
                    'method': 'exact_match',
                    'pdf_idx': pdf_item['original_index'],
                    'docx_idx': docx_item['index'],
                    'pdf_text': pdf_item['original_text'],
                    'docx_text': docx_item['original_text'],
                    'score': 1.0,
                    'reason': 'Точное совпадение нормализованного текста'
                }
                match_log.append(log_entry)
                
                print(f"  ✅ ТОЧНОЕ СОВПАДЕНИЕ: PDF {pdf_item['original_index']+1} -> DOCX {docx_item['index']+1}")
                print(f"      PDF: '{pdf_item['original_text']}'")
                print(f"      DOCX: '{docx_item['original_text']}'")
                break
    
    # МЕТОД 2: Улучшенное Fuzzy matching v3 с токенизацией
    print("\n🔍 МЕТОД 2: Улучшенное Fuzzy matching v3")
    for pdf_item in normalized_pdf:
        if pdf_item['original_index'] in matching:
            continue
            
        best_match = None
        best_score = 0
        
        for docx_item in normalized_docx:
            if docx_item['index'] in used_docx:
                continue
            
            # Используем улучшенное сравнение v3
            content_score = improved_fuzzy_matching_v3(pdf_item['original_text'], docx_item['original_text'])
            
            # Дополнительный бонус за схожесть контекста
            context_score = 0
            if pdf_item['context'] and docx_item['context']:
                context_score = fuzz.partial_ratio(pdf_item['context'], docx_item['context']) / 100.0 * 0.1
            
            total_score = content_score + context_score
            
            print(f"    PDF {pdf_item['original_index']+1} vs DOCX {docx_item['index']+1}: содержание={content_score:.2f}, контекст={context_score:.2f}, итого={total_score:.2f}")
            
            if total_score > best_score and total_score > 0.70:  # ПОНИЖЕН порог для токенизации
                best_score = total_score
                best_match = docx_item
        
        if best_match:
            matching[pdf_item['original_index']] = best_match['index']
            used_docx.add(best_match['index'])
            
            log_entry = {
                'method': 'fuzzy_match',
                'pdf_idx': pdf_item['original_index'],
                'docx_idx': best_match['index'],
                'pdf_text': pdf_item['original_text'],
                'docx_text': best_match['original_text'],
                'score': best_score,
                'reason': f'Fuzzy matching v3 с баллом {best_score:.2f}'
            }
            match_log.append(log_entry)
            
            print(f"  ✅ FUZZY СОВПАДЕНИЕ v3: PDF {pdf_item['original_index']+1} -> DOCX {best_match['index']+1} (балл: {best_score:.2f})")
            print(f"      PDF: '{pdf_item['original_text']}'")
            print(f"      DOCX: '{best_match['original_text']}'")
    
    # МЕТОД 3: Позиционное сопоставление (только при хорошей схожести)
    print("\n📍 МЕТОД 3: Позиционное сопоставление")
    remaining_pdf = [item for item in normalized_pdf if item['original_index'] not in matching]
    remaining_docx = [item for item in normalized_docx if item['index'] not in used_docx]
    
    print(f"Осталось несопоставленных: PDF={len(remaining_pdf)}, DOCX={len(remaining_docx)}")
    
    if len(remaining_pdf) > 0 and len(remaining_docx) > 0 and abs(len(remaining_pdf) - len(remaining_docx)) <= 2:
        remaining_pdf.sort(key=lambda x: x['page'])
        remaining_docx.sort(key=lambda x: x['position'])
        
        min_count = min(len(remaining_pdf), len(remaining_docx))
        
        for i in range(min_count):
            pdf_item = remaining_pdf[i]
            docx_item = remaining_docx[i]
            
            # Проверяем минимальную схожесть
            similarity = improved_fuzzy_matching_v3(pdf_item['original_text'], docx_item['original_text'])
            
            if similarity > 0.3:  # Минимальная схожесть
                matching[pdf_item['original_index']] = docx_item['index']
                used_docx.add(docx_item['index'])
                
                log_entry = {
                    'method': 'position_match',
                    'pdf_idx': pdf_item['original_index'],
                    'docx_idx': docx_item['index'],
                    'pdf_text': pdf_item['original_text'],
                    'docx_text': docx_item['original_text'],
                    'score': 0.5 + similarity * 0.3,
                    'reason': f'Позиционное сопоставление с схожестью {similarity:.2f}'
                }
                match_log.append(log_entry)
                
                print(f"  ✅ ПОЗИЦИОННОЕ СОВПАДЕНИЕ: PDF {pdf_item['original_index']+1} -> DOCX {docx_item['index']+1} (схожесть: {similarity:.2f})")
                print(f"      PDF: '{pdf_item['original_text']}' (стр. {pdf_item['page']})")
                print(f"      DOCX: '{docx_item['original_text']}' (поз. {docx_item['position']})")
            else:
                print(f"  ❌ Позиционное сопоставление отклонено из-за низкой схожести ({similarity:.2f})")
    else:
        print("  ⚠️ Позиционное сопоставление пропущено")
    
    # Логирование несопоставленных формул
    print("\n❌ НЕСОПОСТАВЛЕННЫЕ ФОРМУЛЫ:")
    for pdf_item in normalized_pdf:
        if pdf_item['original_index'] not in matching:
            print(f"  PDF {pdf_item['original_index']+1}: '{pdf_item['original_text']}' (стр. {pdf_item['page']})")
    
    for docx_item in normalized_docx:
        if docx_item['index'] not in used_docx:
            print(f"  DOCX {docx_item['index']+1}: '{docx_item['original_text']}' (поз. {docx_item['position']})")
    
    # Итоговая статистика
    print(f"\n=== ИТОГОВАЯ СТАТИСТИКА v3 ===")
    print(f"Всего сопоставлений: {len(matching)}")
    print(f"Точных совпадений: {len([log for log in match_log if log['method'] == 'exact_match'])}")
    print(f"Fuzzy совпадений: {len([log for log in match_log if log['method'] == 'fuzzy_match'])}")
    print(f"Позиционных совпадений: {len([log for log in match_log if log['method'] == 'position_match'])}")
    
    return matching

# Остальные функции остаются без изменений...
# (Копируем все остальные функции из предыдущей версии)

def remove_existing_comments(docx_path, output_path):
    """БЕЗОПАСНОЕ удаление всех существующих примечаний из DOCX файла"""
    try:
        print(f"Удаление примечаний из {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run_element = run._element
                    
                    comment_elements = []
                    for child in run_element:
                        if child.tag.endswith('}commentRangeStart') or \
                           child.tag.endswith('}commentRangeEnd') or \
                           child.tag.endswith('}commentReference'):
                            comment_elements.append(child)
                    
                    for elem in comment_elements:
                        run_element.remove(elem)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run_element = run._element
                                comment_elements = []
                                for child in run_element:
                                    if child.tag.endswith('}commentRangeStart') or \
                                       child.tag.endswith('}commentRangeEnd') or \
                                       child.tag.endswith('}commentReference'):
                                        comment_elements.append(child)
                                
                                for elem in comment_elements:
                                    run_element.remove(elem)
            
            doc.save(output_path)
            print("✓ Метод 1 (безопасный python-docx) выполнен")
            
            try:
                test_doc = Document(output_path)
                print("✓ Файл прошел проверку целостности")
                return True
            except Exception as e:
                print(f"⚠️ Файл поврежден после метода 1: {e}")
                
        except Exception as e:
            print(f"Ошибка метода 1: {e}")
        
        print("Применяем метод 2 (осторожная работа с ZIP)")
        
        temp_path = output_path + '.temp'
        import shutil
        shutil.copy2(docx_path, temp_path)
        
        with zipfile.ZipFile(temp_path, 'r') as zip_read:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                for item in zip_read.infolist():
                    data = zip_read.read(item.filename)
                    
                    if item.filename == 'word/comments.xml':
                        print("  Удален word/comments.xml")
                        continue
                    
                    if item.filename == 'word/document.xml':
                        try:
                            content = data.decode('utf-8')
                            
                            content = re.sub(r'<w:commentRangeStart[^>]*?/>', '', content)
                            content = re.sub(r'<w:commentRangeEnd[^>]*?/>', '', content)
                            content = re.sub(r'<w:commentReference[^>]*?/>', '', content)
                            
                            data = content.encode('utf-8')
                            print("  Обработан word/document.xml (регулярные выражения)")
                            
                        except Exception as e:
                            print(f"  Ошибка обработки document.xml: {e}")
                    
                    elif item.filename == '[Content_Types].xml':
                        try:
                            content = data.decode('utf-8')
                            lines = content.split('\n')
                            filtered_lines = [line for line in lines if 'comments' not in line.lower()]
                            data = '\n'.join(filtered_lines).encode('utf-8')
                            print("  Обработан [Content_Types].xml")
                        except Exception as e:
                            print(f"  Ошибка обработки [Content_Types].xml: {e}")
                    
                    elif item.filename == 'word/_rels/document.xml.rels':
                        try:
                            content = data.decode('utf-8')
                            lines = content.split('\n')
                            filtered_lines = [line for line in lines if 'comments' not in line.lower()]
                            data = '\n'.join(filtered_lines).encode('utf-8')
                            print("  Обработан word/_rels/document.xml.rels")
                        except Exception as e:
                            print(f"  Ошибка обработки document.xml.rels: {e}")
                    
                    zip_write.writestr(item, data)
        
        os.unlink(temp_path)
        
        try:
            test_doc = Document(output_path)
            print("✓ Метод 2 (ZIP) выполнен и файл прошел проверку")
            return True
        except Exception as e:
            print(f"⚠️ Файл поврежден после метода 2: {e}")
            
    except Exception as e:
        print(f"КРИТИЧЕСКАЯ ОШИБКА при удалении примечаний: {e}")
        try:
            import shutil
            shutil.copy2(docx_path, output_path)
            return False
        except:
            return False

def find_docx_images_with_positions(docx_path):
    """Находит все изображения в DOCX с их позициями и контекстом"""
    try:
        doc = Document(docx_path)
        images_info = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                if run._element.xpath('.//pic:pic'):
                    context_before = ""
                    context_after = ""
                    
                    for i in range(max(0, para_idx - 2), para_idx):
                        context_before += doc.paragraphs[i].text + " "
                    
                    for i in range(run_idx):
                        context_before += paragraph.runs[i].text + " "
                    
                    for i in range(run_idx + 1, len(paragraph.runs)):
                        context_after += paragraph.runs[i].text + " "
                    
                    for i in range(para_idx + 1, min(len(doc.paragraphs), para_idx + 3)):
                        context_after += doc.paragraphs[i].text + " "
                    
                    images_info.append({
                        'paragraph_idx': para_idx,
                        'run_idx': run_idx,
                        'context_before': context_before.strip()[-100:],
                        'context_after': context_after.strip()[:100],
                        'paragraph_text': paragraph.text,
                        'position_in_doc': para_idx
                    })
        
        print(f"Найдено {len(images_info)} изображений в DOCX")
        for i, img in enumerate(images_info):
            print(f"  Изображение {i+1}: параграф {img['paragraph_idx']}, контекст: '{img['context_before'][-50:]}' ... '{img['context_after'][:50]}'")
        
        return images_info
        
    except Exception as e:
        print(f"Ошибка поиска изображений в DOCX: {e}")
        return []

def match_pdf_images_to_docx(pdf_images, docx_images):
    """УЛУЧШЕННОЕ сопоставление изображений PDF с DOCX"""
    matching = {}
    
    if not pdf_images or not docx_images:
        return matching
    
    print(f"Сопоставление: {len(pdf_images)} изображений PDF с {len(docx_images)} изображениями DOCX")
    
    pdf_by_page = {}
    for i, pdf_img in enumerate(pdf_images):
        page = pdf_img.get('page', 1)
        if page not in pdf_by_page:
            pdf_by_page[page] = []
        pdf_by_page[page].append((i, pdf_img))
    
    docx_idx = 0
    for page in sorted(pdf_by_page.keys()):
        page_images = pdf_by_page[page]
        
        for pdf_idx, pdf_img in page_images:
            if docx_idx < len(docx_images):
                matching[pdf_idx] = docx_idx
                print(f"  PDF изображение {pdf_idx} (стр. {page}) -> DOCX изображение {docx_idx}")
                docx_idx += 1
            else:
                print(f"  PDF изображение {pdf_idx} (стр. {page}) -> НЕТ СООТВЕТСТВИЯ в DOCX")
    
    return matching

def calculate_formula_content_similarity(text1, text2):
    """УЛУЧШЕННАЯ функция вычисления схожести содержания формул"""
    if not text1 or not text2:
        return 0.0
    
    norm1 = normalize_formula_text_advanced(text1)
    norm2 = normalize_formula_text_advanced(text2)
    
    if norm1 == norm2:
        return 1.0
    
    clean1 = ''.join(norm1.split())
    clean2 = ''.join(norm2.split())
    
    if clean1 == clean2:
        return 0.95
    
    if clean1 and clean2:
        if clean1 in clean2 or clean2 in clean1:
            shorter = min(len(clean1), len(clean2))
            longer = max(len(clean1), len(clean2))
            if shorter >= 3:
                return (shorter / longer) * 0.8
    
    set1 = set(clean1)
    set2 = set(clean2)
    
    if not set1 or not set2:
        return 0.0
    
    common_chars = set1 & set2
    all_chars = set1 | set2
    
    base_similarity = len(common_chars) / len(all_chars)
    
    math_ops = set('=+-*/^()[]{}')
    math1 = set1 & math_ops
    math2 = set2 & math_ops
    
    if math1 and math2:
        common_math = math1 & math2
        math_bonus = len(common_math) * 0.2
        base_similarity += math_bonus
    
    return min(base_similarity, 1.0)

def calculate_context_similarity(context1, context2):
    """УЛУЧШЕННАЯ функция вычисления схожести контекста"""
    if not context1 or not context2:
        return 0.0
    
    words1 = set(re.findall(r'\b[а-яёa-z]{3,}\b', context1.lower()))
    words2 = set(re.findall(r'\b[а-яёa-z]{3,}\b', context2.lower()))
    
    if not words1 or not words2:
        return 0.0
    
    common_words = words1 & words2
    all_words = words1 | words2
    
    return len(common_words) / len(all_words) if all_words else 0.0

def try_commenting(doc, para_index, comment_text, method="first"):
    """Добавляет примечание к параграфу различными методами"""
    try:
        if para_index >= len(doc.paragraphs):
            return False
        
        para = doc.paragraphs[para_index]
        runs = para.runs
        
        if not runs:
            runs = [para.add_run(" ")]
        
        if method == "first":
            target_runs = runs[0]
        elif method == "last":
            target_runs = runs[-1]
        else:
            target_runs = runs
        
        comment = doc.add_comment(runs=target_runs, text=comment_text, author="ГОСТ Анализатор", initials="ГА")
        print(f"Добавлено примечание методом {method}: {comment_text[:50]}...")
        return True
        
    except Exception as e:
        print(f"Ошибка добавления примечания методом {method}: {e}")
        return False

def annotate_docx_with_issues(docx_path, pdf_path, analysis_results, output_path):
    """ФИНАЛЬНАЯ аннотация DOCX файла с поддержкой формул v3"""
    try:
        doc = Document(docx_path)
        
        docx_images = find_docx_images_with_positions(docx_path)
        docx_formulas = find_docx_formulas_with_positions(docx_path)
        
        from matcher.docx_checks import check_docx_hyphenation, check_docx_double_spaces, check_docx_margins
        
        docx_issues = []
        docx_issues.extend(check_docx_hyphenation(docx_path))
        docx_issues.extend(check_docx_double_spaces(docx_path))
        docx_issues.extend(check_docx_margins(docx_path))
        
        analysis_results['docx_checks'] = docx_issues
        
        if analysis_results.get("images") and docx_images:
            pdf_images = analysis_results["images"]
            image_matching = match_pdf_images_to_docx(pdf_images, docx_images)
            
            print(f"=== ДОБАВЛЕНИЕ ПРИМЕЧАНИЙ К ИЗОБРАЖЕНИЯМ ===")
            
            for pdf_idx, pdf_img in enumerate(pdf_images):
                if not pdf_img.get('gost_compliant', True):
                    issues = []
                    if not pdf_img.get('is_centered', True):
                        issues.append("изображение не центрировано")
                    if not pdf_img.get('margins_ok', True):
                        issues.append("нарушены поля")
                    if not pdf_img.get('has_empty_line', True):
                        issues.append("отсутствует пустая строка перед изображением")
                    
                    if issues:
                        comment_text = f"📷 ИЗОБРАЖЕНИЕ (стр. {pdf_img['page']}): {'; '.join(issues)}"
                        
                        if pdf_idx in image_matching:
                            docx_idx = image_matching[pdf_idx]
                            if docx_idx < len(docx_images):
                                docx_img = docx_images[docx_idx]
                                para_idx = docx_img['paragraph_idx']
                                
                                print(f"Добавляем примечание к изображению {docx_idx+1} в параграфе {para_idx}")
                                try_commenting(doc, para_idx, comment_text, "first")
                            else:
                                print(f"ОШИБКА: docx_idx {docx_idx} вне диапазона")
                        else:
                            print(f"PDF изображение {pdf_idx} не имеет соответствия в DOCX")
        
        if analysis_results.get("formulas") and docx_formulas:
            pdf_formulas = analysis_results["formulas"]
            formula_matching = match_pdf_formulas_to_docx(pdf_formulas, docx_formulas)
            
            print(f"=== ДОБАВЛЕНИЕ ПРИМЕЧАНИЙ К ФОРМУЛАМ ===")
            
            for pdf_idx, pdf_formula in enumerate(pdf_formulas):
                if not pdf_formula.get('gost_compliant', True):
                    issues = []
                    if not pdf_formula.get('is_centered', True):
                        issues.append("формула не центрирована")
                    if not pdf_formula.get('margins_ok', True):
                        issues.append("нарушены поля")
                    if not pdf_formula.get('has_numbering', True):
                        issues.append("отсутствует нумерация")
                    
                    if issues:
                        comment_text = f"🧮 ФОРМУЛА (стр. {pdf_formula['page']}): {'; '.join(issues)} | Текст: '{pdf_formula.get('text', '')[:30]}...'"
                        
                        if pdf_idx in formula_matching:
                            docx_idx = formula_matching[pdf_idx]
                            if docx_idx < len(docx_formulas):
                                docx_formula = docx_formulas[docx_idx]
                                para_idx = docx_formula['paragraph_idx']
                                
                                print(f"Добавляем примечание к формуле {docx_idx+1} в параграфе {para_idx}")
                                try_commenting(doc, para_idx, comment_text, "first")
                            else:
                                print(f"ОШИБКА: docx_idx {docx_idx} вне диапазона")
                        else:
                            print(f"PDF формула {pdf_idx} не имеет соответствия в DOCX")
        
        # Остальная логика для таблиц, текста и т.д. остается без изменений...
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Ошибка при аннотировании DOCX: {e}")
        return False
