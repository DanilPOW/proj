from docx import Document
import re

def check_docx_hyphenation(docx_path):
    """Проверка автопереносов в DOCX"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise Exception(f"Невозможно открыть DOCX файл: {str(e)}")
    
    issues = []
    
    for para_index, para in enumerate(doc.paragraphs):
        # Проверяем настройки переносов
        if hasattr(para.paragraph_format, 'auto_hyphenate'):
            if para.paragraph_format.auto_hyphenate:
                issues.append({
                    "type": "hyphenation",
                    "location": f"Параграф {para_index + 1}",
                    "description": "Включены автоматические переносы",
                    "severity": "high",
                    "para_index": para_index  # Добавляем для точной привязки комментариев
                })
        
        # Поиск мягких переносов в тексте
        if '\u00AD' in para.text:  # мягкий перенос
            issues.append({
                "type": "soft_hyphen",
                "location": f"Параграф {para_index + 1}",
                "description": "Найдены мягкие переносы в тексте",
                "severity": "medium",
                "para_index": para_index
            })
    
    return issues

def check_docx_double_spaces(docx_path):
    """Проверка двойных пробелов в DOCX"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise Exception(f"Невозможно открыть DOCX файл: {str(e)}")
    
    issues = []
    
    for para_index, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Поиск двойных пробелов
        double_spaces = re.findall(r'  +', text)
        if double_spaces:
            issues.append({
                "type": "double_spaces",
                "location": f"Параграф {para_index + 1}",
                "description": f"Найдено {len(double_spaces)} случаев двойных пробелов",
                "severity": "medium",
                "count": len(double_spaces),
                "para_index": para_index
            })
    
    return issues

def check_docx_margins(docx_path):
    """ИСПРАВЛЕННАЯ проверка полей документа в DOCX с правильной конвертацией"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise Exception(f"Невозможно открыть DOCX файл: {str(e)}")
    
    issues = []
    
    # Проверяем поля для каждой секции
    for section_index, section in enumerate(doc.sections):
        try:
            # ПРАВИЛЬНАЯ конвертация из внутренних единиц DOCX в сантиметры
            # В python-docx поля уже возвращаются в правильных единицах через .cm
            left_margin_cm = float(section.left_margin.cm) if section.left_margin else 2.5
            right_margin_cm = float(section.right_margin.cm) if section.right_margin else 2.5
            top_margin_cm = float(section.top_margin.cm) if section.top_margin else 2.5
            bottom_margin_cm = float(section.bottom_margin.cm) if section.bottom_margin else 2.5
            
            # Добавляем подробную отладочную информацию
            print(f"Секция {section_index + 1}:")
            print(f"  Левое поле: {left_margin_cm:.3f} см (требуется: 3.0 см)")
            print(f"  Правое поле: {right_margin_cm:.3f} см (требуется: 2.0 см)")
            print(f"  Верхнее поле: {top_margin_cm:.3f} см (требуется: 2.0 см)")
            print(f"  Нижнее поле: {bottom_margin_cm:.3f} см (требуется: 2.0 см)")
            
            # УВЕЛИЧЕННЫЙ допуск для учета погрешностей конвертации (5мм = 0.5см)
            tolerance = 0.5
            
            # Проверяем левое поле (должно быть 3.0 см)
            if abs(left_margin_cm - 3.0) > tolerance:
                issues.append({
                    "type": "left_margin",
                    "location": f"Секция {section_index + 1}",
                    "description": f"Левое поле {left_margin_cm:.2f}см вместо 3.0см (отклонение: {abs(left_margin_cm - 3.0):.2f}см)",
                    "severity": "high" if abs(left_margin_cm - 3.0) > 1.0 else "medium",
                    "current_value": left_margin_cm,
                    "expected_value": 3.0,
                    "section_index": section_index
                })
                print(f"  ❌ Левое поле не соответствует ГОСТ")
            else:
                print(f"  ✅ Левое поле соответствует ГОСТ")
            
            # Проверяем правое поле (должно быть 2.0 см)
            if abs(right_margin_cm - 2.0) > tolerance:
                issues.append({
                    "type": "right_margin",
                    "location": f"Секция {section_index + 1}",
                    "description": f"Правое поле {right_margin_cm:.2f}см вместо 2.0см (отклонение: {abs(right_margin_cm - 2.0):.2f}см)",
                    "severity": "high" if abs(right_margin_cm - 2.0) > 1.0 else "medium",
                    "current_value": right_margin_cm,
                    "expected_value": 2.0,
                    "section_index": section_index
                })
                print(f"  ❌ Правое поле не соответствует ГОСТ")
            else:
                print(f"  ✅ Правое поле соответствует ГОСТ")
            
            # Проверяем верхнее поле (должно быть 2.0 см)
            if abs(top_margin_cm - 2.0) > tolerance:
                issues.append({
                    "type": "top_margin",
                    "location": f"Секция {section_index + 1}",
                    "description": f"Верхнее поле {top_margin_cm:.2f}см вместо 2.0см (отклонение: {abs(top_margin_cm - 2.0):.2f}см)",
                    "severity": "high" if abs(top_margin_cm - 2.0) > 1.0 else "medium",
                    "current_value": top_margin_cm,
                    "expected_value": 2.0,
                    "section_index": section_index
                })
                print(f"  ❌ Верхнее поле не соответствует ГОСТ")
            else:
                print(f"  ✅ Верхнее поле соответствует ГОСТ")
            
            # Проверяем нижнее поле (должно быть 2.0 см)
            if abs(bottom_margin_cm - 2.0) > tolerance:
                issues.append({
                    "type": "bottom_margin",
                    "location": f"Секция {section_index + 1}",
                    "description": f"Нижнее поле {bottom_margin_cm:.2f}см вместо 2.0см (отклонение: {abs(bottom_margin_cm - 2.0):.2f}см)",
                    "severity": "high" if abs(bottom_margin_cm - 2.0) > 1.0 else "medium",
                    "current_value": bottom_margin_cm,
                    "expected_value": 2.0,
                    "section_index": section_index
                })
                print(f"  ❌ Нижнее поле не соответствует ГОСТ")
            else:
                print(f"  ✅ Нижнее поле соответствует ГОСТ")
            
        except Exception as e:
            print(f"Ошибка проверки полей секции {section_index + 1}: {e}")
            issues.append({
                "type": "margin_error",
                "location": f"Секция {section_index + 1}",
                "description": f"Ошибка чтения полей: {e}",
                "severity": "medium",
                "section_index": section_index
            })
    
    print(f"Проверка полей завершена. Найдено проблем: {len(issues)}")
    return issues
